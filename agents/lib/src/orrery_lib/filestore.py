"""Local-filesystem document store — the replacement for Google Drive.

Files live under ORRERY_FILES_ROOT (default /var/lib/orrery/files). This
module is the agent's read/write surface over that tree:

  - search(): filename + text-content match (PDFs are matched by filename
    only — they are NOT extracted, preserving the finder-not-oracle rule).
  - read():   extract text for Markdown/text/Word; for binaries return a
    pointer to the file's location rather than its bytes.
  - write_text() / save_bytes(): create a NEW file (unique name on
    collision) and make a git commit recording the change.

Versioning: the tree is a git repo; every write commits with the agent's
identity. No Drive API anywhere here.
"""
from __future__ import annotations

import io
import os
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path

FILES_ROOT = Path(os.environ.get("ORRERY_FILES_ROOT", "/var/lib/orrery/files"))
ENGINEERING_ROOT = FILES_ROOT / "engineering"

# Commit identity for agent-initiated writes (override for user-initiated).
GIT_AUTHOR_NAME = os.environ.get("ORRERY_GIT_AUTHOR_NAME", "Orrery Engineering Agent")
GIT_AUTHOR_EMAIL = os.environ.get("ORRERY_GIT_AUTHOR_EMAIL", "agent@orrery.local")

_TEXT_EXT = {".md", ".markdown", ".txt", ".csv"}
_DOCX_EXT = {".docx"}
_PDF_EXT = {".pdf"}
_XLSX_EXT = {".xlsx"}
_ODT_EXT = {".odt"}
_EML_EXT = {".eml"}
_SNIPPET_CHARS = 240

# Skipped when listing/searching a directory (VCS + placeholder noise).
_DIR_SKIP = {".git", ".gitkeep", ".DS_Store"}


@dataclass
class FileHit:
    """One search result. `path` is relative to the search root."""

    path: str
    name: str
    snippet: str = ""


# ── path safety + naming ────────────────────────────────────────────


def _resolve(root: Path, relpath: str) -> Path:
    """Resolve relpath under root, refusing anything that escapes it."""
    p = (root / relpath).resolve()
    root_r = root.resolve()
    if p != root_r and root_r not in p.parents:
        raise ValueError(f"path escapes the store root: {relpath!r}")
    return p


def _safe_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._ -]+", "_", name).strip() or "untitled"
    return cleaned


def _unique_path(directory: Path, filename: str) -> Path:
    target = directory / filename
    if not target.exists():
        return target
    stem, suffix = os.path.splitext(filename)
    i = 1
    while True:
        cand = directory / f"{stem}-{i}{suffix}"
        if not cand.exists():
            return cand
        i += 1


# ── text extraction ─────────────────────────────────────────────────


def _docx_to_text(data: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    lines = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


def _pdf_to_text(data: bytes) -> str:
    import pdfplumber

    out: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    return "\n\n".join(out).strip()


def _xlsx_to_text(data: bytes) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"# {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines).strip()


def _odt_to_text(data: bytes) -> str:
    from odf import teletype
    from odf.opendocument import load

    doc = load(io.BytesIO(data))
    lines: list[str] = []

    # Walk the body in document order, emitting one line per paragraph/heading.
    # (getElementsByType(P)+getElementsByType(H) would float all headings above
    # all paragraphs; a recursive walk preserves the real order, including text
    # nested in tables/lists/frames.)
    def walk(node) -> None:
        for child in getattr(node, "childNodes", []):
            qname = getattr(child, "qname", None)
            if qname is not None and qname[1] in ("p", "h"):
                text = teletype.extractText(child)
                if text.strip():
                    lines.append(text)
            else:
                walk(child)

    walk(doc.text)
    return "\n".join(lines).strip()


def _html_to_text(html: str) -> str:
    """Crude HTML → text: drop script/style, strip tags, unescape entities.
    Good enough for reading an HTML email body; no dependency added."""
    import html as _html

    text = re.sub(r"(?is)<(script|style)\b.*?</\1>", "", html)
    text = re.sub(r"(?is)<(br|/p|/div|/tr|/li)\s*>", "\n", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    text = _html.unescape(text)
    text = re.sub(r"[ \t]+", " ", text)
    return re.sub(r"\n[ \t]+", "\n", text).strip()


def _eml_to_text(data: bytes) -> str:
    """Render a .eml message as readable text: the key headers, then the
    plain-text body (falling back to a stripped HTML body), then a list of
    any attachment filenames."""
    from email import message_from_bytes
    from email.policy import default

    msg = message_from_bytes(data, policy=default)

    header_lines = [
        f"{h}: {msg[h]}" for h in ("From", "To", "Cc", "Date", "Subject") if msg[h]
    ]

    plain: str | None = None
    html: str | None = None
    attachments: list[str] = []
    for part in msg.walk():
        if part.is_multipart():
            continue
        if part.get_content_disposition() == "attachment":
            if part.get_filename():
                attachments.append(part.get_filename())
            continue
        ctype = part.get_content_type()
        if ctype == "text/plain" and plain is None:
            plain = part.get_content()
        elif ctype == "text/html" and html is None:
            html = part.get_content()

    body = plain if plain is not None else (_html_to_text(html) if html else "")

    out = "\n".join(header_lines)
    if attachments:
        out += "\nAttachments: " + ", ".join(attachments)
    return (out + "\n\n" + (body or "").strip()).strip()


def extract_text(path: Path) -> str | None:
    """Deterministic text extraction for catalog/search (Tier 0). Returns None
    for binaries we can't read here (images, archives) and OCR-only PDFs."""
    ext = path.suffix.lower()
    try:
        if ext in _TEXT_EXT:
            return path.read_text(encoding="utf-8", errors="replace")
        if ext in _DOCX_EXT:
            return _docx_to_text(path.read_bytes())
        if ext in _PDF_EXT:
            return _pdf_to_text(path.read_bytes()) or None
        if ext in _XLSX_EXT:
            return _xlsx_to_text(path.read_bytes()) or None
        if ext in _ODT_EXT:
            return _odt_to_text(path.read_bytes()) or None
        if ext in _EML_EXT:
            return _eml_to_text(path.read_bytes()) or None
    except Exception:
        return None  # never let extraction failure block cataloging
    return None


def _snippet(text: str, terms: list[str]) -> str:
    low = text.lower()
    pos = min((low.find(t) for t in terms if t in low), default=-1)
    if pos < 0:
        return text[:_SNIPPET_CHARS].strip()
    start = max(0, pos - 60)
    return text[start : start + _SNIPPET_CHARS].strip().replace("\n", " ")


# ── search / read ───────────────────────────────────────────────────


def search(root: Path, query: str, k: int = 5) -> list[FileHit]:
    """Filename + text-content search. All query terms must appear in the
    filename or the file's extracted text. Filename matches rank higher.
    PDFs/binaries match on filename only."""
    terms = [t for t in query.lower().split() if t]
    if not terms or not root.exists():
        return []
    scored: list[tuple[int, FileHit]] = []
    for p in root.rglob("*"):
        if not p.is_file() or p.name == ".gitkeep":
            continue
        name_l = p.name.lower()
        text = extract_text(p)
        haystack = name_l + "\n" + (text.lower() if text else "")
        if not all(t in haystack for t in terms):
            continue
        in_name = all(t in name_l for t in terms)
        snip = "" if text is None else _snippet(text, terms)
        scored.append(
            (2 if in_name else 1, FileHit(path=str(p.relative_to(root)), name=p.name, snippet=snip))
        )
    scored.sort(key=lambda s: (-s[0], s[1].path))
    return [hit for _, hit in scored[:k]]


def read(root: Path, relpath: str) -> str:
    p = _resolve(root, relpath)
    if not p.is_file():
        raise FileNotFoundError(relpath)
    text = extract_text(p)
    if text is None:
        return (
            f"[{p.name}] is a binary file ({p.suffix or 'no ext'}). It is stored "
            f"at '{relpath}' on the file server — open it directly; its contents "
            "are not extracted."
        )
    return text


def list_dir(root: Path, relpath: str = "") -> list[dict]:
    """List the immediate entries of a directory under `root`. Returns a dict
    per entry with `name`, `type` ('dir' or 'file'), and `path` (relative to
    root). Folders sort first, then files, both alphabetical. Raises
    NotADirectoryError if `relpath` is a file, FileNotFoundError if missing."""
    p = _resolve(root, relpath)
    if not p.exists():
        raise FileNotFoundError(relpath)
    if not p.is_dir():
        raise NotADirectoryError(relpath)
    entries: list[dict] = []
    for child in sorted(p.iterdir(), key=lambda c: (c.is_file(), c.name.lower())):
        if child.name in _DIR_SKIP:
            continue
        entries.append(
            {
                "name": child.name,
                "type": "file" if child.is_file() else "dir",
                "path": str(child.relative_to(root)),
            }
        )
    return entries


# ── writes (with git versioning) ────────────────────────────────────


def git_commit(
    rel_paths: list[str],
    message: str,
    *,
    author_name: str | None = None,
    author_email: str | None = None,
) -> None:
    """Stage paths and commit in the FILES_ROOT repo. Defaults to the agent
    identity; pass author_name/email to attribute a user- or system-initiated
    write."""
    name = author_name or GIT_AUTHOR_NAME
    email = author_email or GIT_AUTHOR_EMAIL
    env = {
        **os.environ,
        "GIT_AUTHOR_NAME": name,
        "GIT_AUTHOR_EMAIL": email,
        "GIT_COMMITTER_NAME": name,
        "GIT_COMMITTER_EMAIL": email,
    }
    base = ["git", "-C", str(FILES_ROOT)]
    subprocess.run([*base, "add", *rel_paths], check=True, env=env, capture_output=True)
    result = subprocess.run(
        [*base, "commit", "-m", message], env=env, capture_output=True, text=True
    )
    if result.returncode != 0 and "nothing to commit" not in (
        result.stdout + result.stderr
    ):
        raise RuntimeError(f"git commit failed: {result.stderr or result.stdout}")


def write_text(rel_dir: str, filename: str, content: str, message: str) -> FileHit:
    """Write a NEW text file under FILES_ROOT/rel_dir, unique on collision,
    then commit. Returns a hit with the path relative to FILES_ROOT."""
    directory = _resolve(FILES_ROOT, rel_dir)
    directory.mkdir(parents=True, exist_ok=True)
    target = _unique_path(directory, _safe_filename(filename))
    target.write_text(content, encoding="utf-8")
    rel = str(target.relative_to(FILES_ROOT))
    git_commit([rel], message)
    return FileHit(path=rel, name=target.name)


def save_bytes(rel_dir: str, filename: str, data: bytes, message: str) -> FileHit:
    """Write a NEW binary file under FILES_ROOT/rel_dir, unique on
    collision, then commit. Returns a hit with the FILES_ROOT-relative path."""
    directory = _resolve(FILES_ROOT, rel_dir)
    directory.mkdir(parents=True, exist_ok=True)
    target = _unique_path(directory, _safe_filename(filename))
    target.write_bytes(data)
    rel = str(target.relative_to(FILES_ROOT))
    git_commit([rel], message)
    return FileHit(path=rel, name=target.name)


# ── the agent's reader ──────────────────────────────────────────────


class FileStoreReader:
    """Read-only search/read/list over one or more roots under the store.

    Paths passed in and returned are relative to FILES_ROOT (e.g.
    'engineering/specs/x.pdf', 'projects/leak-sensor/attachments/m.eml') so
    they are unambiguous across roots and consistent with the rest of the
    system. Access is bounded to the configured roots — the agent cannot read
    outside them (other functions' folders stay invisible)."""

    def __init__(self, roots: list[Path] | Path):
        if isinstance(roots, Path):
            roots = [roots]
        self.base = FILES_ROOT.resolve()
        seen: set[Path] = set()
        self.roots: list[Path] = []
        for r in roots:
            rr = r.resolve()
            if rr not in seen:
                seen.add(rr)
                self.roots.append(rr)

    def _within_roots(self, p: Path) -> bool:
        return any(p == r or r in p.parents for r in self.roots)

    def search(self, query: str, k: int = 5) -> list[FileHit]:
        out: list[FileHit] = []
        seen: set[str] = set()
        for r in self.roots:
            for h in search(r, query, k=k):
                rel = str((r / h.path).relative_to(self.base))
                if rel not in seen:
                    seen.add(rel)
                    out.append(FileHit(path=rel, name=h.name, snippet=h.snippet))
        return out[:k]

    def read(self, path: str) -> str:
        p = _resolve(self.base, path)
        if not self._within_roots(p):
            raise FileNotFoundError(path)
        return read(self.base, path)

    def list_dir(self, path: str = "") -> list[dict]:
        if not path or path in (".", "/"):
            # Top level: present each configured root as a browsable folder.
            return [
                {"name": r.name, "type": "dir", "path": str(r.relative_to(self.base))}
                for r in self.roots
                if r.is_dir()
            ]
        p = _resolve(self.base, path)
        if not self._within_roots(p):
            raise FileNotFoundError(path)
        return list_dir(self.base, path)


def build_engineering_reader() -> FileStoreReader:
    """Default/CLI reader: the engineering reference corpus only."""
    return FileStoreReader(ENGINEERING_ROOT)


def build_reader(folders: list[str]) -> FileStoreReader:
    """Reader over the given FILES_ROOT-relative folders (e.g. a project's
    'projects/<slug>'), always including the engineering corpus. Used per-run
    to scope the agent to the current container's files — where dropped
    documents and emails land — plus its engineering references."""
    roots: list[Path] = []
    for f in folders:
        cleaned = (f or "").strip().strip("/")
        if cleaned:
            roots.append(FILES_ROOT / cleaned)
    roots.append(ENGINEERING_ROOT)
    return FileStoreReader(roots)

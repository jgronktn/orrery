"""Google Drive service-account helpers (agent-agnostic plumbing).

Low-level pieces every Drive-using agent needs: locate the mounted
service-account key, build an authenticated Drive v3 service at a chosen
scope, and convert document bytes to text. Each agent builds its own
reader (search/read scoped to its own folders) on top of these.

The read/write split is enforced by Google's folder ACLs combined with
the scope chosen here: read paths use READONLY_SCOPES, the bounded write
paths use WRITE_SCOPES, and the service account is only Editor on the
agent's `drafts/` folder.
"""
from __future__ import annotations

import io
import os
from pathlib import Path

# Path to the mounted service-account key (read-only bind mount).
SERVICE_ACCOUNT_PATH = Path(
    os.environ.get("ORRERY_GOOGLE_SA_JSON", "/app/secrets/service-account.json")
)

READONLY_SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]
# Write scope; folder ACLs (Editor on drafts/ only) are the real boundary.
WRITE_SCOPES = ["https://www.googleapis.com/auth/drive"]

# Google-native mime types we can export to text.
EXPORT_AS_TEXT = {
    "application/vnd.google-apps.document": "text/plain",
    "application/vnd.google-apps.spreadsheet": "text/csv",
    "application/vnd.google-apps.presentation": "text/plain",
}

# Uploaded Word docs (e.g. templates) — downloaded as bytes and parsed
# to text with python-docx. PDFs are deliberately NOT extracted (per
# CLAUDE.md: the agent surfaces datasheet links, the human reads them).
DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def build_service(scopes: list[str]):
    """Construct a Drive v3 service from the mounted key. Imported
    lazily so this module loads without the Google client dependency."""
    from google.oauth2 import service_account
    from googleapiclient.discovery import build

    creds = service_account.Credentials.from_service_account_file(
        str(SERVICE_ACCOUNT_PATH), scopes=scopes
    )
    # cache_discovery=False avoids a noisy warning + filesystem write.
    return build("drive", "v3", credentials=creds, cache_discovery=False)


def docx_to_text(data: bytes) -> str:
    """Extract plain text from a .docx byte blob (paragraphs + tables).
    python-docx is imported lazily so this module loads without it."""
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    lines = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()

"""Drive access — READ-ONLY reasoning side.

INVARIANT (CLAUDE.md): the agent's reasoning tools are read-only. This
module exposes ONLY search + read, and builds its Google credential with
the read-only scope. The single write capability — creating a new doc in
engineering/drafts/ — lives in the separate `draft.py` module (which
builds its own write-scoped credential and which the agent's reasoning
loop never imports). Do not add a write method here.

The read/write split is ultimately enforced by Google: the service
account is shared Viewer on engineering/ and Editor only on
engineering/drafts/. The code mirrors that intent (read-only scope here)
but does not rely on it alone.

Scope of search: the service account can only see what was shared with
it — the engineering/ tree — so a full-text search across everything it
can access IS a search of engineering/. ENGINEERING_FOLDER_ID is used
only as a configuration sanity signal, not to gate visibility.
"""
from __future__ import annotations

import io
import os
from dataclasses import dataclass
from pathlib import Path
from typing import Protocol

from . import NotConfiguredError

# Path to the mounted service-account key (read-only bind mount).
SERVICE_ACCOUNT_PATH = Path(
    os.environ.get("ORRERY_GOOGLE_SA_JSON", "/app/secrets/service-account.json")
)

# The engineering/ folder id — optional, used for setup sanity only.
ENGINEERING_FOLDER_ID = os.environ.get("ORRERY_ENG_DRIVE_FOLDER_ID", "")

# Read-only scope. The write path (draft.py) uses a different scope.
READONLY_SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]

# Google-native mime types we can export to text.
_EXPORT_AS_TEXT = {
    "application/vnd.google-apps.document": "text/plain",
    "application/vnd.google-apps.spreadsheet": "text/csv",
    "application/vnd.google-apps.presentation": "text/plain",
}

# Uploaded Word docs (e.g. templates) — downloaded as bytes and parsed
# to text with python-docx. PDFs are deliberately NOT extracted (per
# CLAUDE.md: the agent surfaces datasheet links, the human reads them).
_DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

_SETUP_HINT = (
    "Drive access is not configured yet. Save the service-account JSON to "
    "the repo root as service-account.json (gitignored), uncomment its "
    "bind mount in docker-compose.yml, and share engineering/ (Viewer) "
    "with the service-account email."
)


@dataclass
class DriveHit:
    """One Drive search result."""

    file_id: str
    name: str
    mime_type: str
    web_view_link: str
    snippet: str = ""


class DriveReader(Protocol):
    """Read-only contract the agent's Drive tool talks to."""

    def search(self, query: str, k: int = 5) -> list[DriveHit]: ...

    def read(self, file_id: str) -> str: ...


class NullDriveReader:
    """Stand-in until the service account is mounted. Never pretends."""

    def search(self, query: str, k: int = 5) -> list[DriveHit]:
        raise NotConfiguredError(_SETUP_HINT)

    def read(self, file_id: str) -> str:
        raise NotConfiguredError(_SETUP_HINT)


def _docx_to_text(data: bytes) -> str:
    """Extract plain text from a .docx byte blob (paragraphs + tables).
    python-docx is imported lazily so the package loads without it."""
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    lines = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


def _build_service(scopes: list[str]):
    """Construct a Drive v3 service from the mounted key. Imported
    lazily so the package imports fine without the Google client dep."""
    from google.oauth2 import service_account
    from googleapiclient.discovery import build

    creds = service_account.Credentials.from_service_account_file(
        str(SERVICE_ACCOUNT_PATH), scopes=scopes
    )
    # cache_discovery=False avoids a noisy warning + filesystem write.
    return build("drive", "v3", credentials=creds, cache_discovery=False)


class ServiceAccountDriveReader:
    """Read-only Drive access via the engineering service account."""

    def __init__(self):
        self._service = _build_service(READONLY_SCOPES)

    def search(self, query: str, k: int = 5) -> list[DriveHit]:
        # Escape single quotes for the Drive query language.
        safe = query.replace("\\", "\\\\").replace("'", "\\'")
        list_kwargs = dict(
            q=f"fullText contains '{safe}' and trashed = false",
            pageSize=max(1, min(k, 25)),
            fields="files(id, name, mimeType, webViewLink)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        )
        # The engineering corpus lives on a Shared Drive. Shared-drive
        # contents are NOT returned by the default (user) corpus even
        # when shared — you must query the drive corpus explicitly.
        if ENGINEERING_FOLDER_ID.startswith("0A"):
            list_kwargs["corpora"] = "drive"
            list_kwargs["driveId"] = ENGINEERING_FOLDER_ID
        resp = self._service.files().list(**list_kwargs).execute()
        return [
            DriveHit(
                file_id=f["id"],
                name=f.get("name", "?"),
                mime_type=f.get("mimeType", ""),
                web_view_link=f.get("webViewLink", ""),
            )
            for f in resp.get("files", [])
        ]

    def read(self, file_id: str) -> str:
        """Return text for Google-native docs; for binaries (PDFs etc.)
        return a pointer to the link rather than raw bytes — per scope,
        the human reads datasheets/PDFs, the agent surfaces the link."""
        from googleapiclient.http import MediaIoBaseDownload

        meta = (
            self._service.files()
            .get(fileId=file_id, fields="id, name, mimeType, webViewLink",
                 supportsAllDrives=True)
            .execute()
        )
        mime = meta.get("mimeType", "")

        # Uploaded Word doc → download bytes, extract text.
        if mime == _DOCX_MIME:
            buf = io.BytesIO()
            request = self._service.files().get_media(
                fileId=file_id, supportsAllDrives=True
            )
            downloader = MediaIoBaseDownload(buf, request)
            done = False
            while not done:
                _, done = downloader.next_chunk()
            return _docx_to_text(buf.getvalue())

        # Google-native doc → export to text/csv.
        export_mime = _EXPORT_AS_TEXT.get(mime)
        if export_mime is None:
            link = meta.get("webViewLink", "")
            return (
                f"[{meta.get('name', file_id)}] is a non-text file "
                f"({mime}). Open it directly: {link}"
            )
        request = self._service.files().export_media(
            fileId=file_id, mimeType=export_mime
        )
        buf = io.BytesIO()
        downloader = MediaIoBaseDownload(buf, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        return buf.getvalue().decode("utf-8", errors="replace")


def build_drive_reader() -> DriveReader:
    """Live reader if the key is mounted, else the null reader.

    Gated on the key file alone — search relies on the service account's
    visibility, not on the folder id, so it works as soon as the key is
    present and the folder is shared."""
    if SERVICE_ACCOUNT_PATH.exists():
        return ServiceAccountDriveReader()
    return NullDriveReader()
